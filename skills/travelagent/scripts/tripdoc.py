# -*- coding: utf-8 -*-
"""Shared Word builder for travel deliverables.

Provides a Doc wrapper (headings, bullets, tables, real hyperlinks, flags) plus
three assembled builders:

    build_site_guide()       per-city site guide: index by type, day sequencing,
                             cluster writeups, closing flags
    build_editorial_guide()  per-source guide: summary, index, the picks
    build_reservation_tracker()  the action list sorted by deadline

The master itinerary is assembled per trip in templates/build_trip.py because
its day-by-day content is trip-specific; it uses the same Doc primitives.

Formatting invariants (references/deliverables-spec.md): Calibri 10.5, 0.5in
margins on all four sides, navy H1/H2, steel H3, #0563C1 underlined links,
Table Grid with navy header row, red bracketed flags, no em dashes.
"""
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import records as R
import hours as H

NAVY = RGBColor(0x04, 0x1E, 0x42)
STEEL = RGBColor(0x2A, 0x4A, 0x6A)
GREY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xB0, 0x1C, 0x1C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LINK_BLUE = '0563C1'
HDR_FILL = '041E42'

_RT_HYPERLINK = ('http://schemas.openxmlformats.org/officeDocument/2006/'
                 'relationships/hyperlink')

# category and site-type band colours for the index tables
BANDS = {
    'Dining': 'F2E9DA', 'Nightlife': 'EFE4E4',
    'Clothes': 'E9E0F0', 'Shopping': 'FBF4DD',
    'Sights-Sacred': 'E7F0E7', 'Sights-Landmarks': 'E4ECF5',
    'Activities': 'F0EAF5', 'Logistics': 'EDEDED',
    'Temple': 'E7F0E7', 'Shrine': 'E3EEE9', 'Garden': 'E9F2E2',
    'Park': 'E9F2E2', 'Church': 'E7F0E7', 'Mosque': 'E7F0E7',
    'Castle': 'F5E9DC', 'Palace': 'F5E9DC', 'Museum': 'E4ECF5',
    'Landmark': 'EAEFF6', 'District': 'EFEFE6', 'Viewpoint': 'E8F1F6',
    'Observation Deck': 'E8F1F6', 'Memorial': 'F0E8E8', 'Market': 'FBF4DD',
}


def band(key):
    return BANDS.get(key, 'FFFFFF')


# ============================================================ Doc primitives

class Doc(object):
    def __init__(self, title=None, subtitle=None, standfirst=None):
        self.d = Document()
        n = self.d.styles['Normal']
        n.font.name = 'Calibri'
        n.font.size = Pt(10.5)
        for s in self.d.sections:
            s.top_margin = s.bottom_margin = Inches(0.5)
            s.left_margin = s.right_margin = Inches(0.5)
        if title:
            self.title(title, subtitle, standfirst)

    # ---- text -------------------------------------------------------------
    def title(self, text, subtitle=None, standfirst=None):
        p = self.d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text.upper())
        r.font.size = Pt(26)
        r.font.bold = True
        r.font.color.rgb = NAVY
        p.paragraph_format.space_after = Pt(2)
        if subtitle:
            p = self.d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(subtitle)
            r.font.size = Pt(11.5)
            r.font.bold = True
            r.font.color.rgb = STEEL
            p.paragraph_format.space_after = Pt(3)
        if standfirst:
            self.para(standfirst, size=9.5, italic=True, color=GREY, after=6)

    def h(self, text, level=1):
        p = self.d.add_heading(text, level=level)
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.bold = True
            r.font.color.rgb = NAVY if level < 3 else STEEL
            r.font.size = Pt(15 if level == 1 else 12.5 if level == 2 else 11)
        p.paragraph_format.space_before = Pt(9 if level < 3 else 5)
        p.paragraph_format.space_after = Pt(3)
        return p

    def h1(self, t):
        return self.h(t, 1)

    def h2(self, t):
        return self.h(t, 2)

    def h3(self, t):
        return self.h(t, 3)

    def para(self, text='', size=10.5, italic=False, bold=False, color=None,
             after=3, align=None):
        p = self.d.add_paragraph()
        if text:
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.italic = italic
            r.font.bold = bold
            if color is not None:
                r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(after)
        if align is not None:
            p.alignment = align
        return p

    def bullet(self, text, level=1, bold_lead=None, flag=None, size=10.5,
               link=None, link_text='Map'):
        """level 1..3. bold_lead renders before text in bold. flag renders in red."""
        style = 'List Bullet' if level == 1 else 'List Bullet %d' % min(level, 3)
        try:
            p = self.d.add_paragraph(style=style)
        except KeyError:
            p = self.d.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * level)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.space_before = Pt(0)
        if bold_lead:
            r = p.add_run(bold_lead)
            r.font.bold = True
            r.font.size = Pt(size)
            if text:
                sep = p.add_run('  ')
                sep.font.size = Pt(size)
        if text:
            r = p.add_run(text)
            r.font.size = Pt(size)
        if link:
            p.add_run('  ').font.size = Pt(size)
            self.link(p, link, link_text, size=size - 1)
        if flag:
            r = p.add_run('   [%s]' % flag)
            r.font.size = Pt(size - 1.5)
            r.font.bold = True
            r.font.color.rgb = RED
        return p

    def link(self, paragraph, url, text, size=8.5, bold=False):
        part = paragraph.part
        rid = part.relate_to(url, _RT_HYPERLINK, is_external=True)
        el = OxmlElement('w:hyperlink')
        el.set(qn('r:id'), rid)
        run = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), 'Calibri')
        rf.set(qn('w:hAnsi'), 'Calibri')
        rpr.append(rf)
        if bold:
            rpr.append(OxmlElement('w:b'))
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rpr.append(u)
        c = OxmlElement('w:color')
        c.set(qn('w:val'), LINK_BLUE)
        rpr.append(c)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size * 2)))
        rpr.append(sz)
        run.append(rpr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        run.append(t)
        el.append(run)
        paragraph._p.append(el)
        return el

    def page_break(self):
        self.d.add_page_break()

    # ---- tables -----------------------------------------------------------
    @staticmethod
    def shade(cell, hexfill):
        tcpr = cell._tc.get_or_add_tcPr()
        s = OxmlElement('w:shd')
        s.set(qn('w:val'), 'clear')
        s.set(qn('w:color'), 'auto')
        s.set(qn('w:fill'), hexfill)
        tcpr.append(s)

    @staticmethod
    def cell_text(cell, text, size=8, bold=False, color=None, link=None,
                  doc=None):
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if link and doc is not None:
            doc.link(p, link, text, size=size + 0.5, bold=bold)
            return p
        r = p.add_run(text if text is not None else '')
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return p

    def table(self, headers, rows, widths=None, size=8, band_col=None):
        """rows: list of lists. A cell may be a plain string or a
        (text, url) tuple to render as a hyperlink.
        band_col: index of the column whose value picks the row fill colour."""
        t = self.d.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, htxt in enumerate(headers):
            c = t.rows[0].cells[i]
            self.cell_text(c, htxt, size=size + 1, bold=True, color=WHITE)
            self.shade(c, HDR_FILL)
        for row in rows:
            cells = t.add_row().cells
            fill = band(row[band_col]) if band_col is not None and len(row) > band_col else None
            for i, val in enumerate(row):
                if isinstance(val, (tuple, list)) and len(val) == 2:
                    self.cell_text(cells[i], val[0], size=size + 0.5,
                                   link=val[1], doc=self)
                else:
                    self.cell_text(cells[i], '' if val is None else str(val), size=size)
                if fill:
                    self.shade(cells[i], fill)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    if i < len(row.cells):
                        row.cells[i].width = Inches(w)
        return t

    def save(self, path):
        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
        self.d.save(path)
        return path


# ============================================================ shared renderers

def res_short(rec):
    """Compact reservation cell for tables."""
    tag = (rec.get('res') or 'RES:none').replace('RES:', '')
    if tag == 'none':
        return '-'
    how = (rec.get('res_how') or '').strip()
    if how.lower() in ('none', 'n/a', 'na', '-', 'no', 'nothing'):
        how = ''
    if tag == 'walk-in':
        return 'walk-in (%s)' % how if how else 'walk-in'
    return '%s (%s)' % (tag, how) if how else tag


def hours_cell(rec):
    h = H.for_note(rec.get('hours') or 'unknown, verify', max_len=70)
    closed = (rec.get('closed') or '').strip()
    return '%s; %s' % (h, closed) if closed else h


def _flag_run(doc, rec):
    fl = R.flag_list(rec)
    if not fl:
        return
    p = doc.para(after=2)
    r = p.add_run('[' + '] ['.join(fl) + ']')
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RED


def entry_block(doc, rec, show_why=True):
    """One venue or site writeup: name, meta line with Maps link, note, why."""
    p = doc.para(after=1)
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(rec['name'])
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = STEEL
    if rec.get('local_name'):
        r2 = p.add_run('   ' + rec['local_name'])
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = GREY

    meta = doc.para(after=1)
    bits = [rec.get('subcategory', ''), rec.get('neighborhood', ''),
            hours_cell(rec), res_short(rec), rec.get('price', '')]
    r = meta.add_run('  |  '.join(b for b in bits if b) + '   ')
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = GREY
    doc.link(meta, R.maps_url(rec['maps_query'] or rec['name']), 'Map', size=9)

    _flag_run(doc, rec)
    if rec.get('note'):
        doc.para(rec['note'], size=10, after=2)
    if show_why and rec.get('why'):
        doc.para(rec['why'], size=10, after=4)


# ============================================================ site guide

def build_site_guide(out_path, city, recs, sequence=None, intro=None,
                     seq_intro=None, closing=None, subtitle=None):
    """Per-city site guide.

    recs      : records for this city (all categories; sights lead the index)
    sequence  : list of dicts {day, theme, morning, midday, late, evening,
                               closed} rendered as the day-by-day table
    closing   : list of (heading, [bullet, ...]) tuples for the closing flags
    """
    doc = Doc(city,
              subtitle or 'Sites, History, Architecture and Sequencing',
              'Site names link to Google Maps. Hours are in compact notation: '
              'day tokens M Tu W Th F Sa Su, 24-hour times, CLOSED marks the '
              'weekly closure, LO is last order, "last entry" is the museum '
              'cutoff. Verify hours 48 hours out. %d entries.' % len(recs))

    if intro:
        for para in intro:
            doc.para(para, size=10.5, after=4)

    # ---- index by type
    doc.h1('Site Index (by type)')
    doc.para('Grouped by type, then neighborhood, so nearby stops can be '
             'clustered into one day. The Res column is the booking lead time.',
             size=9, italic=True, color=GREY, after=4)
    sights = [r for r in recs if r['category'].startswith('Sights')]
    others = [r for r in recs if not r['category'].startswith('Sights')]
    ordered = (sorted(sights, key=lambda r: (r['subcategory'], r['neighborhood'], r['name']))
               + sorted(others, key=lambda r: (r['category'], r['subcategory'], r['name'])))
    rows = []
    for r in ordered:
        rows.append([r['subcategory'],
                     (r['name'], R.maps_url(r['maps_query'] or r['name'])),
                     r['neighborhood'], hours_cell(r), res_short(r)])
    doc.table(['Type', 'Site', 'Neighborhood', 'Hours', 'Res'], rows,
              widths=[1.15, 2.15, 1.35, 2.55, 1.05], band_col=0)
    doc.page_break()

    # ---- sequencing
    if sequence:
        doc.h1('Suggested Day-by-Day Sequence')
        doc.para(seq_intro or 'Ordered to minimise backtracking and to keep '
                 'every stop open on the day it is scheduled. The Closed today '
                 'column lists anchors shut that day.',
                 size=9, italic=True, color=GREY, after=4)
        rows = []
        for s in sequence:
            rows.append(['%s\n%s' % (s.get('day', ''), s.get('theme', '')),
                         s.get('morning', ''), s.get('midday', ''),
                         s.get('late', ''), s.get('evening', ''),
                         s.get('closed', '')])
        doc.table(['Day / Theme', 'Morning', 'Midday', 'Late afternoon',
                   'Evening', 'Closed today'], rows,
                  widths=[1.15, 1.5, 1.5, 1.5, 1.5, 1.05])
        doc.page_break()

    # ---- detailed guide by cluster
    doc.h1('Detailed Guide')
    for cl in R.clusters(recs, city):
        cr = [r for r in recs if r['cluster'] == cl]
        if not cr:
            continue
        doc.h2(cl)
        for r in sorted(cr, key=lambda x: (x['category'], x['subcategory'], x['name'])):
            entry_block(doc, r)

    # ---- closing flags
    doc.page_break()
    doc.h1('Flags and Booking')
    booking = R.needs_booking(recs)
    if booking:
        doc.h2('Requires booking')
        rows = [[r['name'], r['res'].replace('RES:', ''), r.get('res_how', ''),
                 r['neighborhood'], r.get('price', '')] for r in booking]
        doc.table(['Site', 'Lead time', 'How to book', 'Neighborhood', 'Cost'],
                  rows, widths=[2.0, 0.95, 2.5, 1.4, 1.4])
    closed_notes = [r for r in recs if (r.get('closed') or '').strip()]
    if closed_notes:
        doc.h2('Date-specific and seasonal closures')
        for r in closed_notes:
            doc.bullet(r['closed'], bold_lead=r['name'] + ':')
    flagged = [r for r in recs if R.flag_list(r)]
    if flagged:
        doc.h2('Other flags')
        for r in sorted(flagged, key=lambda x: x['name']):
            doc.bullet('', bold_lead=r['name'] + ':', flag='; '.join(R.flag_list(r)))
    for heading, bullets in (closing or []):
        doc.h2(heading)
        for b in bullets:
            doc.bullet(b)

    return doc.save(out_path)


# ============================================================ editorial guide

def build_editorial_guide(out_path, city, recs, source_label,
                          outline=None, thin_on=None):
    """Per-source-family guide. outline: list of strings summarising coverage."""
    from collections import Counter
    doc = Doc('%s  %s' % (city, source_label),
              '%s Picks  |  Restaurants, Shopping, Sights, Activities' % source_label,
              'Curated from %s. Place names link to Google Maps. Filtered to the '
              'travel dates: no hotels, no lodging, no out-of-season picks. '
              '%d places.' % (source_label, len(recs)))

    doc.h1('Summary')
    cc = Counter(r['category'] for r in recs)
    doc.para('  |  '.join('%s %d' % (c, cc[c]) for c in R.CATEGORIES if cc[c]),
             size=10, bold=True, color=STEEL, after=5)
    for line in (outline or []):
        doc.bullet(line)
    if thin_on:
        doc.para('Thin coverage: ' + thin_on, size=9.5, italic=True,
                 color=GREY, after=4)
    con = R.consensus(recs)
    if con:
        doc.h2('Consensus picks (3 or more independent sources)')
        for r in sorted(con, key=lambda x: (x['category'], x['name'])):
            doc.bullet(r['note'][:150], bold_lead=r['name'],
                       link=R.maps_url(r['maps_query'] or r['name']))

    doc.h1('Index (by category)')
    doc.para('Neighborhood clusters nearby stops. Hours and Res carry the '
             'practical constraints.', size=9, italic=True, color=GREY, after=4)
    ordered = sorted(recs, key=lambda r: (R.CATEGORIES.index(r['category'])
                                          if r['category'] in R.CATEGORIES else 99,
                                          r['subcategory'], r['name']))
    rows = [[r['category'],
             (r['name'], R.maps_url(r['maps_query'] or r['name'])),
             r['subcategory'], r['neighborhood'], hours_cell(r), res_short(r)]
            for r in ordered]
    doc.table(['Category', 'Place', 'Type', 'Neighborhood', 'Hours', 'Res'],
              rows, widths=[1.15, 1.9, 1.2, 1.25, 2.15, 0.85], band_col=0)
    doc.page_break()

    doc.h1('The Picks')
    for cat in R.CATEGORIES:
        cr = [r for r in ordered if r['category'] == cat]
        if not cr:
            continue
        doc.h2(cat)
        for r in cr:
            entry_block(doc, r, show_why=False)

    return doc.save(out_path)


# ============================================================ tracker

def build_reservation_tracker(out_path, recs, trip_label, backups=None):
    """The action list. backups: dict mapping anchor name to a list of
    fallback record dicts, used to fill the 'Backup if it fails' column."""
    doc = Doc('Reservation Tracker', trip_label,
              'Sorted by how far ahead the booking must happen. Everything with '
              'a lead time of a few days or longer appears here. Tick as booked '
              'and record the confirmation.')
    booking = R.needs_booking(recs, cutoff='RES:2-3d')
    if not booking:
        doc.para('Nothing in this record set requires advance booking.',
                 italic=True, color=GREY)
        return doc.save(out_path)

    by_tag = {}
    for r in booking:
        by_tag.setdefault(r['res'], []).append(r)
    for tag in R.RES_TAGS:
        group = by_tag.get(tag)
        if not group:
            continue
        doc.h1('%s  (%d)' % (tag.replace('RES:', 'Book '), len(group)))
        rows = []
        for r in group:
            bk = backups.get(r['name'], []) if backups else []
            bk_txt = ', '.join('%s (%s)' % (b['name'], b['res'].replace('RES:', ''))
                               for b in bk[:2]) or ''
            rows.append([r['name'], r['city'], r.get('res_how', ''),
                         r.get('price', ''), bk_txt, '', ''])
        doc.table(['What', 'City', 'How to book', 'Cost',
                   'Backup if it fails', 'Booked?', 'Confirmation'],
                  rows, widths=[1.8, 0.75, 2.1, 1.3, 2.1, 0.7, 1.0])
    return doc.save(out_path)


if __name__ == '__main__':
    print(__doc__)
